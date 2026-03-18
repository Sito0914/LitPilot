#!/usr/bin/env python3
"""
Literature Processing Pipeline
Processes PDFs and Word documents in the incoming/ folder using the Claude API
with BMAD agent definitions for structured summarisation.

Features:
  - PyMuPDF (fitz) for high-quality PDF text extraction
  - Concurrent API calls for faster batch processing
  - Retry with exponential backoff for transient API errors
  - Duplicate detection via processing log
  - Automatic chunking for very long papers

Outputs:
  - Detailed markdown note per paper → summaries/individual/
  - One Excel row per paper → summaries/literature_tracker.xlsx
  - Processing log → processing_log.csv
"""

import csv
import os
import sys
import time
import hashlib
from datetime import datetime
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, as_completed

# Allow importing config from the project root
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from dotenv import load_dotenv
load_dotenv(Path(__file__).resolve().parent.parent / ".env")

import anthropic
import fitz  # PyMuPDF — much better text extraction than PyPDF2
from docx import Document as DocxDocument

from config import (
    PROJECT_ROOT, INCOMING_DIR, PROCESSED_DIR, SUMMARIES_DIR,
    AGENT_DIR, LOG_FILE, EXCEL_TRACKER, TRACKER_HEADERS,
    ANALYSIS_MODEL, ANALYSIS_MAX_TOKENS, MAX_PDF_PAGES,
    MAX_RETRIES, MAX_CONCURRENT, CHARS_PER_TOKEN,
    ANALYSIS_INPUT_COST, ANALYSIS_OUTPUT_COST, CONTEXT_LIMIT,
    SUPPORTED_EXTENSIONS,
)


# ══════════════════════════════════════════════
# DUPLICATE DETECTION
# ══════════════════════════════════════════════

def get_file_hash(file_path: Path) -> str:
    """Generate MD5 hash of a file for duplicate detection."""
    hasher = hashlib.md5()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            hasher.update(chunk)
    return hasher.hexdigest()


def load_processed_hashes() -> set:
    """Load hashes of previously processed files from the log."""
    hashes = set()
    if not LOG_FILE.exists():
        return hashes
    with open(LOG_FILE, "r", newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            if row.get("status") == "SUCCESS" and row.get("hash"):
                hashes.add(row["hash"])
    return hashes


def log_processing(filename: str, status: str, output_path: str, file_hash: str = ""):
    """Append to the processing log."""
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    headers = ["timestamp", "filename", "status", "output", "hash"]

    write_header = not LOG_FILE.exists()

    with open(LOG_FILE, "a", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        if write_header:
            writer.writerow(headers)
        writer.writerow([timestamp, filename, status, output_path, file_hash])


# ══════════════════════════════════════════════
# TOKEN ESTIMATION
# ══════════════════════════════════════════════

def estimate_tokens(text: str) -> int:
    """Estimate token count from character length."""
    return len(text) // CHARS_PER_TOKEN


def estimate_cost(input_tokens: int, output_tokens: int = ANALYSIS_MAX_TOKENS) -> float:
    """Estimate API cost in USD for a single call."""
    input_cost = (input_tokens / 1_000_000) * ANALYSIS_INPUT_COST
    output_cost = (output_tokens / 1_000_000) * ANALYSIS_OUTPUT_COST
    return input_cost + output_cost


# ══════════════════════════════════════════════
# TEXT EXTRACTION
# ══════════════════════════════════════════════

def _tesseract_available() -> bool:
    """Check whether Tesseract OCR is installed on the system."""
    import shutil
    return shutil.which("tesseract") is not None


def extract_text_from_pdf(pdf_path: Path) -> str:
    """Extract text from a PDF using PyMuPDF (fitz) for better quality.
    Falls back to per-page OCR (via Tesseract) for scanned pages."""
    try:
        doc = fitz.open(str(pdf_path))
        total_pages = len(doc)

        if total_pages > MAX_PDF_PAGES:
            print(f"  WARNING: {pdf_path.name} has {total_pages} pages. "
                  f"Processing first {MAX_PDF_PAGES} pages only.")

        pages_to_process = min(total_pages, MAX_PDF_PAGES)
        text_parts = []
        text_page_count = 0
        ocr_page_count = 0
        has_tesseract = _tesseract_available()

        for i in range(pages_to_process):
            page = doc[i]
            page_text = page.get_text("text")

            if len(page_text.strip()) > 200:
                # Enough text for a real page (not just headers/footers/captions)
                text_parts.append(f"--- Page {i+1} ---\n{page_text}")
                text_page_count += 1
            else:
                # Page has little or no text — attempt OCR
                if not has_tesseract:
                    continue
                try:
                    tp = page.get_textpage_ocr(flags=0, language="eng", dpi=300)
                    ocr_text = page.get_text("text", textpage=tp)
                    if ocr_text.strip():
                        text_parts.append(f"--- Page {i+1} (OCR) ---\n{ocr_text}")
                        ocr_page_count += 1
                except Exception:
                    pass

        doc.close()
        full_text = "\n\n".join(text_parts)

        # Print status based on the four cases
        if ocr_page_count == 0 and text_page_count > 0:
            # Case 1: Normal PDF
            pass
        elif text_page_count > 0 and ocr_page_count > 0:
            # Case 2: Mixed
            print(f"  Mixed PDF: {text_page_count} text pages, "
                  f"{ocr_page_count} OCR pages")
        elif ocr_page_count > 0 and text_page_count == 0:
            # Case 3: All OCR
            print(f"  Scanned PDF: all {ocr_page_count} pages processed via OCR")
        elif text_page_count == 0 and ocr_page_count == 0:
            # Case 4: No text and no OCR
            if not has_tesseract:
                print(f"  WARNING: No text extracted from {pdf_path.name}. "
                      f"This appears to be a scanned PDF but Tesseract is not installed.")
                print(f"  Install it with: brew install tesseract (macOS) "
                      f"or apt install tesseract-ocr (Linux)")
            else:
                print(f"  WARNING: Very little text extracted from {pdf_path.name}, "
                      f"even with OCR.")

        return full_text

    except Exception as e:
        print(f"  ERROR extracting text from {pdf_path.name}: {e}")
        return ""


def extract_text_from_docx(docx_path: Path) -> str:
    """Extract text from a Word document, including tables."""
    try:
        doc = DocxDocument(str(docx_path))
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)

        full_text = "\n\n".join(text_parts)

        if len(full_text.strip()) < 100:
            print(f"  WARNING: Very little text extracted from {docx_path.name}.")

        return full_text

    except Exception as e:
        print(f"  ERROR extracting text from {docx_path.name}: {e}")
        return ""


def extract_text(file_path: Path) -> str:
    """Route to the correct extractor based on file type."""
    ext = file_path.suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    else:
        print(f"  Unsupported file type: {ext}")
        return ""


# ══════════════════════════════════════════════
# CLAUDE API (with retry logic)
# ══════════════════════════════════════════════

def load_agent_prompt(agent_name: str) -> str:
    """Load a BMAD agent definition."""
    agent_file = AGENT_DIR / f"{agent_name}.md"
    if not agent_file.exists():
        print(f"ERROR: Agent file not found: {agent_file}")
        sys.exit(1)
    return agent_file.read_text(encoding="utf-8")


def call_claude_with_retry(client, model, max_tokens, system, messages) -> tuple:
    """Call Claude API with exponential backoff retry on transient errors.
    Returns (response_text, input_tokens, output_tokens)."""
    for attempt in range(1, MAX_RETRIES + 1):
        try:
            response = client.messages.create(
                model=model,
                max_tokens=max_tokens,
                system=system,
                messages=messages,
            )
            result = ""
            for block in response.content:
                if block.type == "text":
                    result += block.text
            return result, response.usage.input_tokens, response.usage.output_tokens

        except anthropic.RateLimitError:
            wait = 2 ** attempt
            print(f"  Rate limited. Waiting {wait}s (attempt {attempt}/{MAX_RETRIES})...")
            time.sleep(wait)

        except anthropic.APIStatusError as e:
            if e.status_code >= 500:
                wait = 2 ** attempt
                print(f"  Server error ({e.status_code}). Waiting {wait}s "
                      f"(attempt {attempt}/{MAX_RETRIES})...")
                time.sleep(wait)
            else:
                raise

        except anthropic.APIConnectionError:
            wait = 2 ** attempt
            print(f"  Connection error. Waiting {wait}s "
                  f"(attempt {attempt}/{MAX_RETRIES})...")
            time.sleep(wait)

    return "ERROR: API request failed after all retries.", 0, 0


def send_to_claude(paper_text: str, agent_prompt: str) -> str:
    """Send paper text to Claude API with the agent prompt."""
    client = anthropic.Anthropic()

    try:
        result, in_tokens, out_tokens = call_claude_with_retry(
            client=client,
            model=ANALYSIS_MODEL,
            max_tokens=ANALYSIS_MAX_TOKENS,
            system=agent_prompt,
            messages=[
                {
                    "role": "user",
                    "content": (
                        "Please analyse the following academic paper and "
                        "produce a structured summary according to your "
                        "instructions.\n\n"
                        f"PAPER TEXT:\n\n{paper_text}"
                    )
                }
            ],
        )
        actual_cost = (
            (in_tokens / 1_000_000) * ANALYSIS_INPUT_COST
            + (out_tokens / 1_000_000) * ANALYSIS_OUTPUT_COST
        )
        print(f"  Actual: {in_tokens:,} input + {out_tokens:,} output tokens (${actual_cost:.2f})")
        return result

    except anthropic.BadRequestError as e:
        if "too long" in str(e).lower() or "token" in str(e).lower():
            print(f"  Paper too long for single request. Chunking...")
            return process_long_paper(paper_text, agent_prompt)
        else:
            print(f"  API ERROR: {e}")
            return f"ERROR: API request failed — {e}"

    except Exception as e:
        print(f"  API ERROR: {e}")
        return f"ERROR: API request failed — {e}"


def process_long_paper(paper_text: str, agent_prompt: str) -> str:
    """Handle papers that exceed the context window by chunking."""
    client = anthropic.Anthropic()

    chunk_chars = 80000 * CHARS_PER_TOKEN
    chunks = [paper_text[i:i + chunk_chars] for i in range(0, len(paper_text), chunk_chars)]

    print(f"  Split into {len(chunks)} chunks...")

    chunk_summaries = []
    for idx, chunk in enumerate(chunks):
        print(f"  Processing chunk {idx + 1}/{len(chunks)}...")
        chunk_result, _, _ = call_claude_with_retry(
            client=client,
            model=ANALYSIS_MODEL,
            max_tokens=2048,
            system=(
                "You are reading part of an academic paper. "
                "Extract: key arguments, methodology details, "
                "findings, and any theoretical frameworks mentioned. "
                "Be thorough but concise."
            ),
            messages=[
                {
                    "role": "user",
                    "content": (
                        f"This is chunk {idx+1} of {len(chunks)} "
                        f"from an academic paper.\n\n{chunk}"
                    )
                }
            ],
        )
        chunk_summaries.append(chunk_result)

    combined = "\n\n---\n\n".join(chunk_summaries)

    result, in_tokens, out_tokens = call_claude_with_retry(
        client=client,
        model=ANALYSIS_MODEL,
        max_tokens=ANALYSIS_MAX_TOKENS,
        system=agent_prompt,
        messages=[
            {
                "role": "user",
                "content": (
                    "Below are summaries of different sections of an "
                    "academic paper (it was too long to process at once). "
                    "Please synthesise these into a single structured "
                    "summary according to your instructions.\n\n"
                    f"SECTION SUMMARIES:\n\n{combined}"
                )
            }
        ],
    )
    actual_cost = (
        (in_tokens / 1_000_000) * ANALYSIS_INPUT_COST
        + (out_tokens / 1_000_000) * ANALYSIS_OUTPUT_COST
    )
    print(f"  Chunked paper actual: {in_tokens:,} input + {out_tokens:,} output tokens (${actual_cost:.2f})")
    return result


# ══════════════════════════════════════════════
# OUTPUT
# ══════════════════════════════════════════════

def save_summary(filename: str, summary: str) -> Path:
    """
    Split Claude's output into:
    1. A detailed markdown note → summaries/individual/
    2. A row in the Excel tracker → summaries/literature_tracker.xlsx
    """
    from openpyxl import Workbook, load_workbook

    CSV_SEPARATOR = "===CSV_DATA==="

    if CSV_SEPARATOR in summary:
        parts = summary.split(CSV_SEPARATOR, 1)
        markdown_part = parts[0].strip()
        csv_line = parts[1].strip()
    else:
        markdown_part = summary
        csv_line = None

    # Save the markdown note
    summary_name = Path(filename).stem + "_note.md"
    output_path = SUMMARIES_DIR / summary_name
    output_path.write_text(markdown_part, encoding="utf-8")

    # Append to Excel tracker
    if csv_line:
        fields = csv_line.split("|")
        # The prompt outputs 13 fields (12 original + DOI).
        # We insert Source File before DOI to keep tracker column order:
        # [1-12 original fields] + [Source File] + [DOI]
        if len(fields) >= 13:
            doi = fields[12].strip()
            fields = fields[:12]
            fields.append(filename)
            fields.append(doi)
        else:
            fields = fields[:12]
            fields.append(filename)
            fields.append("")

        if EXCEL_TRACKER.exists():
            wb = load_workbook(EXCEL_TRACKER)
            ws = wb.active
        else:
            wb = Workbook()
            ws = wb.active
            ws.title = "Literature Tracker"
            ws.append(TRACKER_HEADERS)

        ws.append(fields)
        wb.save(EXCEL_TRACKER)

    return output_path


# ══════════════════════════════════════════════
# SINGLE PAPER PROCESSOR (for concurrent use)
# ══════════════════════════════════════════════

def process_single_paper(file_path: Path, agent_prompt: str, file_hash: str,
                         pre_extracted_text: str = None) -> dict:
    """Process one paper end-to-end. Returns a result dict."""
    result = {
        "file": file_path,
        "status": "FAILED",
        "output": "",
        "hash": file_hash,
    }

    try:
        text = pre_extracted_text if pre_extracted_text is not None else extract_text(file_path)

        if not text or len(text.strip()) < 100:
            result["status"] = "SKIPPED"
            result["output"] = "insufficient text"
            print(f"  [{file_path.name}] SKIPPED: insufficient text")
            return result

        tokens = estimate_tokens(text)
        cost = estimate_cost(tokens)
        print(f"  [{file_path.name}] Extracted {len(text):,} chars (~{tokens:,} tokens, ~${cost:.2f})")

        summary = send_to_claude(text, agent_prompt)

        output_path = save_summary(file_path.name, summary)
        print(f"  [{file_path.name}] Note saved: {output_path.name}")

        result["status"] = "SUCCESS"
        result["output"] = str(output_path)
        return result

    except Exception as e:
        print(f"  [{file_path.name}] FAILED: {e}")
        result["output"] = str(e)
        return result


# ══════════════════════════════════════════════
# MAIN PIPELINE
# ══════════════════════════════════════════════

def process_incoming():
    """Process all PDFs and Word docs in the incoming folder."""

    # Find all supported files
    files = []
    for ext in SUPPORTED_EXTENSIONS:
        files.extend(INCOMING_DIR.glob(f"*{ext}"))
    files.sort(key=lambda f: f.name)

    if not files:
        print("No PDF or DOCX files found in incoming/ folder.")
        print(f"Drop your files into: {INCOMING_DIR}")
        return

    # Duplicate detection
    processed_hashes = load_processed_hashes()
    new_files = []
    for f in files:
        fhash = get_file_hash(f)
        if fhash in processed_hashes:
            print(f"  SKIPPED (duplicate): {f.name}")
            dest = PROCESSED_DIR / f.name
            if dest.exists():
                ts = datetime.now().strftime("%Y%m%d_%H%M%S")
                dest = PROCESSED_DIR / f"{f.stem}_{ts}{f.suffix}"
            f.rename(dest)
        else:
            new_files.append((f, fhash))

    if not new_files:
        print("All files have already been processed. No new papers to analyse.")
        return

    print(f"Found {len(new_files)} new file(s) to process.\n")
    print(f"  PDFs:  {sum(1 for f, _ in new_files if f.suffix.lower() == '.pdf')}")
    print(f"  DOCX:  {sum(1 for f, _ in new_files if f.suffix.lower() == '.docx')}")
    if len(files) > len(new_files):
        print(f"  Skipped {len(files) - len(new_files)} duplicate(s).")

    # Batch cost estimate (extract text first for accurate estimate)
    print("  Estimating costs (extracting text)...")
    extracted_texts = {}
    est_batch_cost = 0.0
    for f, fhash in new_files:
        text = extract_text(f)
        extracted_texts[f] = text
        tokens = estimate_tokens(text)
        est_batch_cost += estimate_cost(tokens)
    print(f"  Est. batch cost: ~${est_batch_cost:.2f}")
    print()

    # Load the analyst agent prompt
    agent_prompt = load_agent_prompt("analyst")

    # Validate that the user has customised the agent prompt
    if "[YOUR TOPIC" in agent_prompt or "[Section Title]" in agent_prompt:
        print("ERROR: Your analyst agent prompt still contains placeholder text.")
        print("Edit bmad/agents/analyst.md before processing papers.")
        print("See the Customisation Guide in README.md.")
        sys.exit(1)

    # Dry-run mode: show what would be processed without calling the API
    if "--dry-run" in sys.argv:
        print("DRY RUN — no API calls will be made.\n")
        for file_path, fhash in new_files:
            text = extracted_texts.get(file_path, "")
            tokens = estimate_tokens(text)
            cost = estimate_cost(tokens)
            print(f"  {file_path.name}: {len(text):,} chars, ~{tokens:,} tokens, ~${cost:.2f}")
        total = sum(
            estimate_cost(estimate_tokens(extracted_texts.get(f, "")))
            for f, _ in new_files
        )
        print(f"\n  Total estimated cost: ~${total:.2f}")
        return

    # Process papers concurrently
    results = []
    with ThreadPoolExecutor(max_workers=MAX_CONCURRENT) as executor:
        future_to_file = {}
        for file_path, fhash in new_files:
            pre_text = extracted_texts.get(file_path)
            future = executor.submit(
                process_single_paper, file_path, agent_prompt, fhash, pre_text
            )
            future_to_file[future] = (file_path, fhash)

        for future in as_completed(future_to_file):
            file_path, fhash = future_to_file[future]
            result = future.result()
            results.append(result)

            if result["status"] in ("SUCCESS", "SKIPPED"):
                destination = PROCESSED_DIR / file_path.name
                if destination.exists():
                    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
                    destination = PROCESSED_DIR / f"{file_path.stem}_{ts}{file_path.suffix}"
                if file_path.exists():
                    file_path.rename(destination)

            log_processing(
                file_path.name, result["status"], result["output"], fhash
            )

    # Summary
    succeeded = sum(1 for r in results if r["status"] == "SUCCESS")
    skipped = sum(1 for r in results if r["status"] == "SKIPPED")
    failed = sum(1 for r in results if r["status"] == "FAILED")

    print()
    print("=" * 60)
    print(f"Done. Processed {len(results)} file(s).")
    print(f"  Succeeded: {succeeded}")
    if skipped:
        print(f"  Skipped:   {skipped}")
    if failed:
        print(f"  Failed:    {failed}")
    print(f"  Notes:     {SUMMARIES_DIR}")
    print(f"  Tracker:   {EXCEL_TRACKER}")
    print(f"  Log:       {LOG_FILE}")


if __name__ == "__main__":
    if not os.environ.get("ANTHROPIC_API_KEY"):
        print("ERROR: ANTHROPIC_API_KEY not found.")
        print(f"Check your .env file at: {PROJECT_ROOT / '.env'}")
        print("Copy .env.example to .env and add your key.")
        sys.exit(1)

    print(f"Project:  {PROJECT_ROOT}")
    print(f"Incoming: {INCOMING_DIR}")
    print()

    process_incoming()
